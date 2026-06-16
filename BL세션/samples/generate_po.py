"""
구매발주서(PO) DOCX 5개 생성 스크립트 (python-docx 필요)

samples/의 BL 샘플 5개(BL-2026-L001 / F001 / F002 / Z001 / B001)와 1:1 대응되는
Purchase Order DOCX를 생성한다.
흐름: PO 발행(이 스크립트) -> 공급사 생산/선적 -> BL 발행(generate_samples.py)

실행: python generate_po.py
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

DARK_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
DARK_BLUE_HEX = "1E3A5F"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

BUYER_NAME = "YeongFab Fashion Co., Ltd."
BUYER_ADDR = "45 Yeonsu-gu, Incheon, Korea"

DEFAULT_PORT_LOADING = "Shanghai, China"
DEFAULT_PORT_DISCHARGE = "Incheon, Korea"

TERMS_AND_CONDITIONS = [
    "Goods must conform to the agreed specifications and approved samples; any deviation requires prior written approval from the Buyer.",
    "Delivery delay beyond the agreed shipment date may incur a penalty of 0.5% of the order value per day, up to a maximum of 5%.",
    "Payment Terms: 30% deposit by T/T upon order confirmation, 70% balance by T/T against a copy of the Bill of Lading.",
    "All shipments must be accompanied by a Commercial Invoice, Packing List, and Certificate of Origin.",
    "This Purchase Order is governed by the laws of the Republic of Korea. Disputes shall be settled by good-faith negotiation, failing which by arbitration in Seoul, Korea.",
    "This document may be amended only by mutual written agreement between Buyer and Vendor prior to shipment.",
]


def _shade(hex_color: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    return shd


def _set_cell(cell, text, bold=False, color=None, size=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align


def make_po(
    filename,
    po_no,
    po_date,
    bl_ref,
    vendor_name,
    vendor_addr,
    items,
    port_of_loading=DEFAULT_PORT_LOADING,
    port_of_discharge=DEFAULT_PORT_DISCHARGE,
    incoterm=None,
):
    incoterm = incoterm or f"FOB {port_of_loading}"

    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PURCHASE ORDER")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    # ── Header info table ────────────────────────────────────────
    info = doc.add_table(rows=5, cols=4)
    info.style = "Table Grid"
    rows_data = [
        ("PO No.:", po_no, "PO Date:", po_date),
        ("Vendor:", vendor_name, "Buyer:", BUYER_NAME),
        ("Vendor Address:", vendor_addr, "Buyer Address:", BUYER_ADDR),
        ("Incoterm:", incoterm, "Related BL No.:", bl_ref),
        ("Port of Loading:", port_of_loading, "Port of Discharge:", port_of_discharge),
    ]
    for r, (l1, v1, l2, v2) in enumerate(rows_data):
        _set_cell(info.cell(r, 0), l1, bold=True)
        _set_cell(info.cell(r, 1), v1)
        _set_cell(info.cell(r, 2), l2, bold=True)
        _set_cell(info.cell(r, 3), v2)

    doc.add_paragraph()

    # ── Item table ────────────────────────────────────────────────
    h = doc.add_paragraph()
    run = h.add_run("ITEMS")
    run.bold = True
    run.font.size = Pt(12)

    cols = ["No.", "Material", "Specification", "Qty", "Unit", "Unit Price (USD)", "Amount (USD)"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, c in enumerate(cols):
        cell = table.cell(0, i)
        _set_cell(cell, c, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell._tc.get_or_add_tcPr().append(_shade(DARK_BLUE_HEX))

    total_amount = 0.0
    for idx, it in enumerate(items, 1):
        amount = it["qty"] * it["unit_price"]
        total_amount += amount
        row = table.add_row()
        values = [
            str(idx),
            it["material"],
            it["spec"],
            f"{it['qty']:,}",
            it["unit"],
            f"{it['unit_price']:,.2f}",
            f"{amount:,.2f}",
        ]
        for i, v in enumerate(values):
            if i in (0, 3, 4):
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif i in (5, 6):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                align = None
            _set_cell(row.cells[i], v, align=align)

    total_row = table.add_row()
    for i in range(5):
        _set_cell(total_row.cells[i], "")
    _set_cell(total_row.cells[5], "TOTAL", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell(total_row.cells[6], f"{total_amount:,.2f}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()

    # ── Terms and conditions ─────────────────────────────────────
    h2 = doc.add_paragraph()
    run = h2.add_run("TERMS AND CONDITIONS")
    run.bold = True
    run.font.size = Pt(12)

    for i, term in enumerate(TERMS_AND_CONDITIONS, 1):
        p = doc.add_paragraph(f"{i}. {term}")
        p.runs[0].font.size = Pt(9)

    doc.add_paragraph()

    sig = doc.add_paragraph(
        "Authorized Signature (Buyer): ______________________     Date: ____________"
    )
    sig.runs[0].font.size = Pt(9)

    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"  Generated: {filename}")


def main():
    print("Generating Purchase Order DOCX files...")

    # 1. 라벨회사 -> Guangzhou Label Supply Co., Ltd. (BL-2026-L001)
    make_po(
        "PO_라벨회사_2026-05-15.docx",
        po_no="PO-2026-L001",
        po_date="2026-05-15",
        bl_ref="BL-2026-L001",
        vendor_name="Guangzhou Label Supply Co., Ltd.",
        vendor_addr="No. 88 Tianhe Industrial Zone, Guangzhou, China",
        items=[
            {"material": "Label Fabric", "spec": "Woven Label Fabric Polyester Satin 25mm Width", "qty": 50000, "unit": "m", "unit_price": 0.04},
            {"material": "Printing Ink", "spec": "Heat Transfer Printing Ink Black 5L per Can", "qty": 80, "unit": "cans", "unit_price": 22.00},
        ],
    )

    # 2. 옷감회사 -> Shanghai Textile Manufacturing Co., Ltd. (BL-2026-F001)
    make_po(
        "PO_옷감회사_2026-05-15.docx",
        po_no="PO-2026-F001",
        po_date="2026-05-15",
        bl_ref="BL-2026-F001",
        vendor_name="Shanghai Textile Manufacturing Co., Ltd.",
        vendor_addr="No. 168 Pudong Industrial Ave, Shanghai, China",
        items=[
            {"material": "Cotton Yarn", "spec": "100% Cotton Ring-Spun Yarn 32s Count", "qty": 500, "unit": "kg", "unit_price": 3.20},
            {"material": "Polyester Yarn", "spec": "Polyester Filament Yarn 150D/48F", "qty": 300, "unit": "kg", "unit_price": 1.80},
        ],
    )

    # 3. 옷감회사 -> Qingdao Premium Textile Co., Ltd. (BL-2026-F002)
    make_po(
        "PO_옷감회사_2026-05-16.docx",
        po_no="PO-2026-F002",
        po_date="2026-05-16",
        bl_ref="BL-2026-F002",
        vendor_name="Qingdao Premium Textile Co., Ltd.",
        vendor_addr="No. 25 Laoshan District, Qingdao, China",
        items=[
            {"material": "Linen Yarn", "spec": "100% Linen Wet-Spun Yarn 24s Count", "qty": 200, "unit": "kg", "unit_price": 6.50},
            {"material": "Wool Yarn", "spec": "100% Merino Wool Yarn 2/48s Count", "qty": 150, "unit": "kg", "unit_price": 9.00},
            {"material": "Mixed Yarn", "spec": "Cotton Polyester Blended Yarn 65/35 32s", "qty": 250, "unit": "kg", "unit_price": 2.50},
        ],
    )

    # 4. 지퍼단추회사 -> YKK Shanghai Trading Co., Ltd. (BL-2026-Z001)
    make_po(
        "PO_지퍼단추회사_2026-05-15.docx",
        po_no="PO-2026-Z001",
        po_date="2026-05-15",
        bl_ref="BL-2026-Z001",
        vendor_name="YKK Shanghai Trading Co., Ltd.",
        vendor_addr="No. 500 Songjiang Industrial Park, Shanghai, China",
        items=[
            {"material": "Zipper Tape", "spec": "Nylon Zipper Tape 25mm Width Assorted Colors", "qty": 2000, "unit": "m", "unit_price": 0.15},
            {"material": "Metal Material", "spec": "Iron Zinc Alloy Ingot for Button Stamping", "qty": 500, "unit": "kg", "unit_price": 4.00},
        ],
    )

    # 5. 지퍼단추회사 -> Wenzhou Button Manufacturing Co., Ltd. (BL-2026-B001)
    make_po(
        "PO_지퍼단추회사_2026-05-16.docx",
        po_no="PO-2026-B001",
        po_date="2026-05-16",
        bl_ref="BL-2026-B001",
        vendor_name="Wenzhou Button Manufacturing Co., Ltd.",
        vendor_addr="No. 12 Ouhai Industrial Zone, Wenzhou, China",
        items=[
            {"material": "Wood Material", "spec": "Natural Hardwood Block Beech for Button Cutting", "qty": 300, "unit": "kg", "unit_price": 2.20},
            {"material": "Plastic Material", "spec": "ABS Plastic Resin Pellets for Button Molding", "qty": 400, "unit": "kg", "unit_price": 1.50},
        ],
    )

    print("Done! 5 PO DOCX files created.")


if __name__ == "__main__":
    main()
